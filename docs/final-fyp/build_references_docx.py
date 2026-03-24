#!/usr/bin/env python3
"""Build REFERENCES-MERGED-HARVARD.docx from REFERENCES-MERGED-HARVARD.txt.

*Text in asterisks* in the .txt becomes italic in Word (Pansilu-style: journal / conference / book).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DIR = Path(__file__).resolve().parent
OUT = DIR / "REFERENCES-MERGED-HARVARD.docx"
SRC = DIR / "REFERENCES-MERGED-HARVARD.txt"


def load_reference_lines(path: Path) -> list[str]:
    raw = path.read_text(encoding="utf-8")
    chunks = raw.split("\n---\n")
    if len(chunks) < 2:
        raise SystemExit(f"Expected at least one '---' block in {path}")
    # First block after initial --- is the reference list (before Notes ---)
    ref_block = chunks[1]
    lines: list[str] = []
    for ln in ref_block.splitlines():
        s = ln.strip()
        if not s:
            continue
        lines.append(s)
    return lines


def add_runs_with_asterisk_italic(paragraph, text: str, font_pt: float = 11) -> None:
    """*wrapped* segments become italic; URLs must not contain lone *."""
    pos = 0
    for m in re.finditer(r"\*([^*]+)\*", text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            r.font.size = Pt(font_pt)
        ri = paragraph.add_run(m.group(1))
        ri.font.size = Pt(font_pt)
        ri.italic = True
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(font_pt)


NOTE = (
    "Notes for Word:\n\n"
    "1. Plain-text source: REFERENCES-MERGED-HARVARD.txt — *asterisks* mark italic (journal / conference / book).\n"
    "2. This .docx applies those italics; you can re-run: python3 build_references_docx.py\n"
    "3. Disambiguate in-text: Yu, R. (2021) vs Yu, X. (2024, 2025).\n"
    "4. Add [Accessed: date] if required.\n"
)


def main() -> None:
    lines = load_reference_lines(SRC)
    if not lines:
        raise SystemExit("No reference lines parsed from .txt")

    doc = Document()
    title = doc.add_heading(
        "References (merged — Harvard, Pansilu-style italics)", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph(
        "Journal / conference / book titles are italic below. Article titles are roman. "
        "Source: REFERENCES-MERGED-HARVARD.txt (*…* markers)."
    )
    p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    for line in lines:
        para = doc.add_paragraph()
        add_runs_with_asterisk_italic(para, line)
        para.paragraph_format.space_after = Pt(8)

    doc.add_page_break()
    doc.add_heading("Notes", level=1)
    for line in NOTE.strip().split("\n"):
        doc.add_paragraph(line, style="List Bullet")

    doc.save(OUT)
    print(f"Wrote {OUT} ({len(lines)} references)")


if __name__ == "__main__":
    main()
