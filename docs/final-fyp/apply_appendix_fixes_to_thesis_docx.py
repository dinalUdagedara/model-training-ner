#!/usr/bin/env python3
"""
Apply appendix letter fixes to w1998730_FYP.docx (body + tables).

- Technical 'Appendix A' -> 'Appendix G' (datasets, notebooks, credentials).
- Full functional test tables: 'Appendix B' -> 'Appendix D' (only the Ch8 sentence).
- Does NOT rename headings 'Appendix A – Survey Results'.

Default thesis path: ../../../crackInt/FINAL/w1998730_FYP.docx (FYP folder = repo’s parent; override with THESIS_DOCX=...).

Creates a .bak copy next to the docx, then overwrites the docx.
"""
from __future__ import annotations

import os
import shutil
from pathlib import Path

from docx import Document

# Ordered: longest / most specific first to avoid partial collisions.
REPLACEMENTS: list[tuple[str, str]] = [
    (
        "may be placed in Appendix B; this section summarizes representative cases",
        "may be placed in Appendix D; this section summarizes representative cases",
    ),
    (
        "Quantitative results here must remain internally consistent with Chapter 07, Appendix A, and the Abstract.",
        "Quantitative results here must remain internally consistent with Chapter 07, Appendix G, and the Abstract.",
    ),
    (
        "see Appendix A for the filename and any version note",
        "see Appendix G for the filename and any version note",
    ),
    (
        "see Appendix A for the exact filename used in the frozen run",
        "see Appendix G for the exact filename used in the frozen run",
    ),
    (
        "see Appendix A for merge filename and provenance",
        "see Appendix G for merge filename and provenance",
    ),
    (
        "trained in the job-poster NER training notebook (Appendix A),",
        "trained in the job-poster NER training notebook (Appendix G),",
    ),
    (
        "listed in Appendix A.",
        "listed in Appendix G.",
    ),
    (
        "frozen run (Appendix A);",
        "frozen run (Appendix G);",
    ),
    (
        "training notebooks (Appendix A);",
        "training notebooks (Appendix G);",
    ),
    (
        "Chapter 07 and Appendix A (frozen training run)",
        "Chapter 07 and Appendix G (frozen training run)",
    ),
    (
        "credentials (Appendix A)",
        "credentials (Appendix G)",
    ),
]


def replace_in_text(text: str) -> tuple[str, int]:
    n = 0
    for old, new in REPLACEMENTS:
        if old in text:
            c = text.count(old)
            text = text.replace(old, new)
            n += c
    return text, n


def patch_document(doc: Document) -> int:
    total = 0
    for p in doc.paragraphs:
        new, n = replace_in_text(p.text)
        if n:
            p.text = new
            total += n
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    new, n = replace_in_text(p.text)
                    if n:
                        p.text = new
                        total += n
    return total


def main() -> None:
    default = Path(__file__).resolve().parents[3] / "crackInt" / "FINAL" / "w1998730_FYP.docx"
    path = Path(os.environ.get("THESIS_DOCX", str(default)))
    if not path.exists():
        raise SystemExit(f"Missing thesis docx: {path}")

    backup = path.with_suffix(".docx.bak-appendix")
    shutil.copy2(path, backup)
    print(f"Backup: {backup}")

    doc = Document(str(path))
    n = patch_document(doc)
    doc.save(str(path))
    print(f"Patched: {path}")
    print(f"Replacement operations applied (substring occurrences): {n}")
    print("Next: open Word, update TOC, insert Appendix D/G/H content from appendix-*-PASTE-INTO-WORD.txt")


if __name__ == "__main__":
    main()
